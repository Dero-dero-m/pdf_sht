import io

import pytest
from docx import Document as DocxDocument
from httpx import AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.models import Document


@pytest.mark.asyncio
async def test_download_docx_returns_word_doc(client: AsyncClient, session: AsyncSession) -> None:
    doc = Document(
        filename="report.pdf",
        page_count=1,
        byte_size=1,
        content_markdown="# Hello\n\nWorld.",
    )
    session.add(doc)
    await session.commit()
    await session.refresh(doc)

    resp = await client.get(f"/documents/{doc.id}/docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "report.docx" in resp.headers["content-disposition"]

    parsed = DocxDocument(io.BytesIO(resp.content))
    assert any(p.text == "Hello" and p.style.name == "Heading 1" for p in parsed.paragraphs)
    assert any(p.text == "World." for p in parsed.paragraphs)


@pytest.mark.asyncio
async def test_download_docx_404(client: AsyncClient) -> None:
    resp = await client.get("/documents/9999/docx")
    assert resp.status_code == 404
