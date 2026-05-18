import io

from docx import Document as DocxDocument

from app.services.docx_export import markdown_to_docx_bytes


def _read(b: bytes) -> "DocxDocument":
    return DocxDocument(io.BytesIO(b))


def test_heading_levels_map_to_word_styles() -> None:
    md = "# H1\n\n## H2\n\n### H3\n\nBody paragraph."
    doc = _read(markdown_to_docx_bytes(md))
    styles = [p.style.name for p in doc.paragraphs if p.text]
    assert styles[:3] == ["Heading 1", "Heading 2", "Heading 3"]


def test_bullet_and_numbered_lists() -> None:
    md = "- one\n- two\n\n1. first\n2. second"
    doc = _read(markdown_to_docx_bytes(md))
    texts = [p.text for p in doc.paragraphs if p.text]
    assert "one" in texts and "two" in texts and "first" in texts and "second" in texts
    list_styles = [p.style.name for p in doc.paragraphs if p.text in ("one", "two")]
    assert all("Bullet" in s or "List" in s for s in list_styles)
