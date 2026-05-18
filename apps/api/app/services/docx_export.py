"""Tiny Markdown→DOCX renderer.

Covers the shapes Claude reliably emits for PDF extraction:
ATX headings (# .. ######), paragraphs, ordered/unordered lists.
Inline formatting is rendered as plain text.
"""
import io
import re

from docx import Document as DocxDocument

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_UL_RE = re.compile(r"^[-*+]\s+(.*)$")
_OL_RE = re.compile(r"^\d+\.\s+(.*)$")


def markdown_to_docx_bytes(markdown: str) -> bytes:
    doc = DocxDocument()
    para_buf: list[str] = []

    def flush_paragraph() -> None:
        if para_buf:
            doc.add_paragraph(" ".join(para_buf).strip())
            para_buf.clear()

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            flush_paragraph()
            continue
        m = _HEADING_RE.match(line)
        if m:
            flush_paragraph()
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            continue
        m = _UL_RE.match(line)
        if m:
            flush_paragraph()
            doc.add_paragraph(m.group(1).strip(), style="List Bullet")
            continue
        m = _OL_RE.match(line)
        if m:
            flush_paragraph()
            doc.add_paragraph(m.group(1).strip(), style="List Number")
            continue
        para_buf.append(line.strip())
    flush_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
